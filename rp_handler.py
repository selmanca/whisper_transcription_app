import os
import base64
import subprocess
import tempfile
import wave
import contextlib
import webrtcvad
import torch
import librosa
import re
from transformers import WhisperProcessor, WhisperForConditionalGeneration
from docx import Document
import runpod
from openai import OpenAI

# Temporary directory
tmpdir = tempfile.gettempdir()

# ----- Audio Processing Functions -----

def convert_audio_to_wav(input_path, output_path):
    command = [
        "ffmpeg", "-y", "-i", input_path,
        "-vn", "-acodec", "pcm_s16le",
        "-ar", "16000", "-ac", "1", "-f", "wav", output_path
    ]
    try:
        subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"FFmpeg conversion failed: {e.stderr.decode('utf-8')}")

def read_wave(path):
    with contextlib.closing(wave.open(path, 'rb')) as wf:
        assert wf.getframerate() == 16000, "Sample rate must be 16kHz"
        assert wf.getnchannels() == 1, "Audio must be mono"
        assert wf.getsampwidth() == 2, "Audio must be 16-bit"
        pcm_data = wf.readframes(wf.getnframes())
    return pcm_data, 16000

def write_wave(path, audio_bytes, sample_rate=16000):
    with wave.open(path, 'wb') as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)
        wf.setframerate(sample_rate)
        wf.writeframes(audio_bytes)

def vad_chunks_to_flexible_chunks(wav_path, aggressiveness=2):
    FRAME_MS, MIN_SEC, MAX_SEC = 30, 25, 30
    audio_bytes, sr = read_wave(wav_path)
    vad = webrtcvad.Vad(aggressiveness)
    frame_size = int(sr * (FRAME_MS / 1000.0) * 2)
    frames = [audio_bytes[i:i+frame_size] for i in range(0, len(audio_bytes), frame_size)]
    # Mark speech frames
    speech_frames = [frame if len(frame) >= frame_size and vad.is_speech(frame, sr) else None for frame in frames]
    chunks, current, sec, start_time, chunk_infos = [], bytearray(), 0, 0, []
    for i, f in enumerate(speech_frames):
        if f is not None:
            current.extend(f)
            sec += FRAME_MS / 1000
        elif current:
            current.extend(b'\x00' * frame_size)
            sec += FRAME_MS / 1000
        if sec >= MAX_SEC or (sec >= MIN_SEC and (i+1 == len(speech_frames) or speech_frames[i+1] is None)):
            end_time = start_time + sec
            start_sample = max(0, int(start_time*sr)*2 - int(1.2*sr)*2)
            end_sample = min(len(audio_bytes), int(end_time*sr)*2 + int(1.2*sr)*2)
            chunks.append(audio_bytes[start_sample:end_sample])
            chunk_infos.append((start_time, end_time))
            current.clear(); start_time = end_time; sec = 0
    if current:
        end_time = start_time + sec
        chunks.append(bytes(current))
        chunk_infos.append((start_time, end_time))
    # Save chunks to disk
    paths = []
    for i, (chunk, (start, end)) in enumerate(zip(chunks, chunk_infos)):
        path = os.path.join(tmpdir, f"vad_chunk_{i}.wav")
        write_wave(path, chunk, sr)
        print(f"Chunk {i}: {start:.2f}s to {end:.2f}s -> {path}")
        paths.append(path)
    return paths

def merge_transcriptions(transcriptions, overlap_words=15):
    if not transcriptions: 
        return ""
    merged = transcriptions[0].strip()
    for current in transcriptions[1:]:
        current = current.strip()
        m_words, c_words = merged.split(), current.split()
        max_ov = 0
        for size in range(overlap_words, 0, -1):
            if m_words[-size:] == c_words[:size]:
                max_ov = size; break
        merged += " " + " ".join(c_words[max_ov:])
    return merged

def add_markdown_paragraph(doc, markdown_text):
    """Parse Markdown and add text to DOCX with **bold** formatting."""
    paragraph = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', markdown_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

# ----- Helper: Process Audio File -----

def process_audio_file(audio_file_path, processor, model, device):
    processed_audio = os.path.join(tmpdir, f"processed_{os.path.basename(audio_file_path)}.wav")
    convert_audio_to_wav(audio_file_path, processed_audio)
    chunk_paths = vad_chunks_to_flexible_chunks(processed_audio)
    transcriptions = []
    for chunk in chunk_paths:
        waveform, _ = librosa.load(chunk, sr=16000)
        input_feats = processor(waveform, sampling_rate=16000, return_tensors="pt"
                                 ).input_features.to(device)
        pred_ids = model.generate(input_feats)
        transcription = processor.batch_decode(pred_ids, skip_special_tokens=True)[0]
        transcriptions.append(transcription.strip())
    return merge_transcriptions(transcriptions)

# ----- RunPod Handler Function -----

def handler(event):
    job_input = event.get("input", {}) or {}
    # Get input as either base64 or file paths (support single or list)
    audio_b64 = job_input.get("audio_base64") or []
    audio_path = job_input.get("audio") or []
    if isinstance(audio_b64, str):
        audio_b64 = [audio_b64]
    if isinstance(audio_path, str):
        audio_path = [audio_path]

    if not (audio_b64 or audio_path):
        return {"error": "No audio input provided."}
    
    # Initialize model components once
    model_name = os.getenv("WHISPER_MODEL_NAME")
    processor = WhisperProcessor.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = WhisperForConditionalGeneration.from_pretrained(model_name).to(device)
    
    all_transcriptions = []
    
    # Process base64 files: decode, write to disk, then process.
    for idx, audio_str in enumerate(audio_b64):
        orig_path = os.path.join(tmpdir, f"input_audio_original_{idx}")
        try:
            if audio_str.startswith("data:"):
                audio_str = audio_str.split(",", 1)[1]
            with open(orig_path, "wb") as f:
                f.write(base64.b64decode(audio_str))
            transcript = process_audio_file(orig_path, processor, model, device)
            all_transcriptions.append(transcript)
        except Exception as e:
            return {"error": f"Error processing base64 audio index {idx}: {e}"}
    
    # Process direct file path audio.
    for idx, path in enumerate(audio_path):
        try:
            transcript = process_audio_file(path, processor, model, device)
            all_transcriptions.append(transcript)
        except Exception as e:
            return {"error": f"Error processing audio file {path}: {e}"}
    
    combined_transcription = "\n\n".join(all_transcriptions)
    
    # LLM Processing using OpenAI API template.
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return {"error": "OPENAI_API_KEY environment variable must be set."}
    client = OpenAI(api_key=api_key)
    system_prompt = (
        "You are ChatGPT. You fix spelling, grammar, punctuation, and formatting errors in Turkish text in radiology report setting.\n"
        "Rules:\n"
        "Fix all errors but do not change the meaning.\n"
        "Use Markdown **bold** for headings or names (e.g., **Heading:**).\n"
        "Add a line break after every sentence.\n"
        "Don't add or remove any major content.\n"
        "Only output the corrected text — no explanations."
    )
    user_message = f'"{combined_transcription}"'
    try:
        response = client.responses.create(
            model="gpt-5",
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.8
        )
        llm_output = response.output_text
    except Exception as e:
        return {"error": f"LLM processing failed: {e}"}
    
    # Create DOCX with LLM-corrected text parsed via Markdown.
    doc = Document()
    add_markdown_paragraph(doc, llm_output.strip())
    output_docx = os.path.join(tmpdir, "transcription.docx")
    doc.save(output_docx)
    try:
        with open(output_docx, "rb") as f:
            docx_b64 = base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        return {"error": f"Could not encode DOCX file: {e}"}
    
    return {"result": docx_b64}

if __name__ == "__main__":
    runpod.serverless.start({"handler": handler})

