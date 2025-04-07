import os
import datetime
import zipfile
import tempfile

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
import torchaudio
from docx import Document
from pydub import AudioSegment
import torch
from transformers import WhisperProcessor, WhisperForConditionalGeneration

app = FastAPI()

# Ensure CUDA is available
if not torch.cuda.is_available():
    raise EnvironmentError("CUDA is not available. Make sure you have a GPU and correct drivers installed.")

# Load the fine-tuned Whisper model from Hugging Face
device = "cuda" if torch.cuda.is_available() else "cpu"
model = WhisperForConditionalGeneration.from_pretrained("sarumanca/whisper-finetuned").to(device)
processor = WhisperProcessor.from_pretrained("sarumanca/whisper-finetuned")

def process_audio_file(file_path: str, language: str = "tr") -> str | None:
    """
    Processes a single audio file:
    - Converts it to WAV (resampling to 16 kHz and converting to mono if needed)
    - Loads and transcribes the audio using the fine-tuned Whisper model
    - Creates a Word (.docx) file with the transcription
    Returns the path to the generated Word file.
    """
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_doc_name = f"{base_name}_{timestamp}.docx"
    output_doc_path = os.path.join(tempfile.gettempdir(), output_doc_name)

    # Convert to WAV if necessary
    try:
        audio = AudioSegment.from_file(file_path)
        if audio.frame_rate != 16000:
            print(f"Resampling from {audio.frame_rate} Hz to 16000 Hz.")
            audio = audio.set_frame_rate(16000)
        if audio.channels != 1:
            print(f"Converting from {audio.channels} channels to mono.")
            audio = audio.set_channels(1)

        temp_audio_path = os.path.join(tempfile.gettempdir(), f"{base_name}_temp.wav")
        audio.export(temp_audio_path, format="wav")
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return None

    # Load audio with torchaudio
    try:
        waveform, sr = torchaudio.load(temp_audio_path)
        if sr != 16000:
            print(f"Resampling inside torchaudio: {sr} Hz to 16000 Hz")
            resampler = torchaudio.transforms.Resample(orig_freq=sr, new_freq=16000)
            waveform = resampler(waveform)
    except Exception as e:
        print(f"Error loading audio with torchaudio: {e}")
        return None

    # Prepare input for the model and transcribe
    try:
        # Ensure waveform is a 1D numpy array
        input_features = processor(waveform.squeeze(), sampling_rate=16000, return_tensors="pt").input_features.to(device)
        predicted_ids = model.generate(input_features)
        transcription = processor.batch_decode(predicted_ids, skip_special_tokens=True)[0]
        print(f"Transcription: {transcription}")
    except Exception as e:
        print(f"Transcription error: {e}")
        transcription = ""

    # Create a Word document with the transcription
    doc = Document()
    doc.add_heading(f"Transcript for {base_name}", level=1)
    doc.add_paragraph("")
    if not transcription.strip():
        doc.add_paragraph("⚠️ No speech was detected or transcription failed.")
    else:
        doc.add_paragraph(transcription)

    doc.save(output_doc_path)

    # Clean up temporary audio file
    if os.path.exists(temp_audio_path):
        os.remove(temp_audio_path)

    return output_doc_path

@app.post("/transcribe/")
async def transcribe_audio(files: list[UploadFile] = File(...), language: str = "tr"):
    """
    Accepts one or more audio files and a language code.
    Processes each file, generating a Word document for each transcription.
    Returns a ZIP archive of the generated Word files.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded.")

    output_files = []
    # Create a temporary directory to hold uploaded and output files
    with tempfile.TemporaryDirectory() as temp_dir:
        for upload in files:
            # Save the uploaded file to the temporary directory
            file_path = os.path.join(temp_dir, upload.filename)
            with open(file_path, "wb") as f:
                f.write(await upload.read())
            doc_path = process_audio_file(file_path, language=language)
            if doc_path:
                output_files.append(doc_path)

        # Create a zip archive of the generated Word files
        zip_path = os.path.join(temp_dir, "transcripts.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))

        # Return the zip archive as the response
        return FileResponse(zip_path, media_type="application/zip", filename="transcripts.zip")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=3000, reload=True)
